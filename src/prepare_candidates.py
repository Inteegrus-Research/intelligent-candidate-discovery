#!/usr/bin/env python3
"""
Prepare raw candidates and JD into clean, compact artifacts.

Inputs:
- data/raw/candidates.jsonl (or .jsonl.gz)
- data/raw/job_description.docx

Outputs:
- data/artifacts/candidate_ids.npy
- data/artifacts/candidate_metadata.parquet
- data/artifacts/candidate_preliminary_features.parquet
- data/artifacts/candidate_text.parquet
"""

import argparse, gzip, json, re
from datetime import date, datetime
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from tqdm import tqdm

SERVICE_COMPANIES = {
    "tcs", "infosys", "wipro", "accenture", "cognizant", "capgemini",
    "mindtree", "hcl", "tech mahindra", "ltimindtree", "l&t infotech",
    "deloitte", "pwc", "ey", "kpmg"
}

SKILL_CANON = {
    "sentence transformers": "sentence-transformers",
    "sentence-transformer": "sentence-transformers",
    "vector db": "vector database",
    "vector database": "vector database",
    "vector databases": "vector database",
    "llms": "llm",
    "large language model": "llm",
    "large language models": "llm",
    "fine tuning llms": "fine-tuning llms",
    "fine-tuning llms": "fine-tuning llms",
    "learning to rank": "learning-to-rank",
    "learning-to-rank": "learning-to-rank",
    "recommendation system": "recommendation systems",
    "recommendation systems": "recommendation systems",
}

def parse_date(s):
    if not s:
        return None
    try:
        return datetime.strptime(str(s), "%Y-%m-%d").date()
    except Exception:
        return None

def days_ago(d, today):
    return (today - d).days if d else None

def clean_text(x):
    if not x:
        return ""
    x = str(x).strip().lower()
    x = re.sub(r"\s+", " ", x)
    return x

def normalize_skill(name):
    name = clean_text(name)
    name = name.replace("&", " and ")
    return SKILL_CANON.get(name, name)

def load_docx_text(path):
    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" ".join(cells))
    return "\n".join(parts)

def compress_career_description(desc, max_chars=150):
    if not desc:
        return ""
    desc = desc.strip()
    first_sent = re.split(r'(?<=[.!?])\s+', desc)[0]
    if len(first_sent) > max_chars:
        return first_sent[:max_chars] + "..."
    return first_sent

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--data", default="data/raw/candidates.jsonl")
    ap.add_argument("--jd", default="data/raw/job_description.docx")
    ap.add_argument("--outdir", default="data/artifacts")
    ap.add_argument("--as_of", default="2026-06-20", help="YYYY-MM-DD for date normalization")
    args = ap.parse_args()

    data_path = Path(args.data)
    jd_path = Path(args.jd)
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    today = datetime.strptime(args.as_of, "%Y-%m-%d").date()

    jd_raw = load_docx_text(jd_path)
    jd_text = clean_text(jd_raw)

    if str(data_path).endswith(".gz"):
        f = gzip.open(data_path, "rt", encoding="utf-8")
    else:
        f = open(data_path, "r", encoding="utf-8")

    ids = []
    meta_rows = []
    feat_rows = []
    text_rows = []

    for line in tqdm(f, total=100000, desc="Processing candidates"):
        c = json.loads(line.strip())
        cid = c["candidate_id"]
        ids.append(cid)

        p = c["profile"]
        sig = c["redrob_signals"]

        signup = parse_date(sig.get("signup_date"))
        last_active = parse_date(sig.get("last_active_date"))

        skills = c.get("skills", [])
        norm_skills = []
        for s in skills:
            norm_skills.append({
                "name": normalize_skill(s.get("name", "")),
                "proficiency": str(s.get("proficiency", "")).lower(),
                "endorsements": int(s.get("endorsements", 0) or 0),
                "duration_months": int(s.get("duration_months", 0) or 0),
            })

        career = c.get("career_history", [])
        career_months = 0
        service_count = 0
        career_titles = []
        for r in career:
            dur = int(r.get("duration_months", 0) or 0)
            career_months += dur
            comp = clean_text(r.get("company", ""))
            if any(sc in comp for sc in SERVICE_COMPANIES):
                service_count += 1
            career_titles.append(clean_text(r.get("title", "")))

        yoe = float(p.get("years_of_experience", 0) or 0)
        career_years = round(career_months / 12.0, 2)
        current_title_clean = clean_text(p.get("current_title", ""))

        title_match_count = sum(1 for t in career_titles if t == current_title_clean)
        title_career_match_ratio = title_match_count / len(career_titles) if career_titles else 0.0

        profile_doc = clean_text(f"{p.get('headline','')}. {p.get('summary','')}")
        skills_doc_compact = ", ".join(s["name"] for s in norm_skills)
        career_parts = []
        for r in career:
            title = clean_text(r.get("title", ""))
            company = clean_text(r.get("company", ""))
            desc = compress_career_description(r.get("description", ""))
            career_parts.append(f"{title} at {company}: {desc}")
        career_doc_compact = " ; ".join(career_parts)
        full_doc = f"{profile_doc}. Skills: {skills_doc_compact}. {career_doc_compact}"

        meta_rows.append({
            "candidate_id": cid,
            "anonymized_name": clean_text(p.get("anonymized_name","")),
            "headline": clean_text(p.get("headline","")),
            "current_title": current_title_clean,
            "current_company": clean_text(p.get("current_company","")),
            "current_industry": clean_text(p.get("current_industry","")),
            "location": clean_text(p.get("location","")),
            "country": clean_text(p.get("country","")),
            "years_of_experience": yoe,
            "career_years_from_history": career_years,
            "days_since_signup": days_ago(signup, today),
            "days_since_active": days_ago(last_active, today),
            "open_to_work_flag": 1 if sig.get("open_to_work_flag") else 0,
            "notice_period_days": int(sig.get("notice_period_days", 90) or 90),
            "verified_email": 1 if sig.get("verified_email") else 0,
            "verified_phone": 1 if sig.get("verified_phone") else 0,
            "linkedin_connected": 1 if sig.get("linkedin_connected") else 0,
            "preferred_work_mode": clean_text(sig.get("preferred_work_mode","")),
            "willing_to_relocate": 1 if sig.get("willing_to_relocate") else 0,
            "github_activity_score": float(sig.get("github_activity_score", -1) or -1),
            "recruiter_response_rate": float(sig.get("recruiter_response_rate", 0) or 0),
            "interview_completion_rate": float(sig.get("interview_completion_rate", 0) or 0),
            "offer_acceptance_rate": float(sig.get("offer_acceptance_rate", -1) or -1),
            "profile_completeness_score": float(sig.get("profile_completeness_score", 0) or 0),
            "profile_views_received_30d": int(sig.get("profile_views_received_30d", 0) or 0),
            "connection_count": int(sig.get("connection_count", 0) or 0),
            "endorsements_received": int(sig.get("endorsements_received", 0) or 0),
            "service_company_count": service_count,
            "career_history_count": len(career),
            "education_count": len(c.get("education", [])),
            "skills_count": len(norm_skills),
            "title_career_match_ratio": title_career_match_ratio,
        })

        advanced_cnt = sum(1 for s in norm_skills if s["proficiency"] in ("advanced", "expert"))
        expert_cnt = sum(1 for s in norm_skills if s["proficiency"] == "expert")
        weighted_skill = 0.0
        for s in norm_skills:
            prof_map = {"beginner":0.25, "intermediate":0.5, "advanced":0.8, "expert":1.0}
            ps = prof_map.get(s["proficiency"], 0.0)
            weighted_skill += ps * np.log1p(s["endorsements"]) * np.log1p(s["duration_months"])

        github_val = float(sig.get("github_activity_score", -1) or -1)
        offer_acc = float(sig.get("offer_acceptance_rate", -1) or -1)

        feat_rows.append({
            "candidate_id": cid,
            "yoe": yoe,
            "career_years_from_history": career_years,
            "days_since_active": days_ago(last_active, today),
            "days_since_signup": days_ago(signup, today),
            "notice_period_days": int(sig.get("notice_period_days", 90) or 90),
            "profile_completeness": float(sig.get("profile_completeness_score", 0) or 0),
            "response_rate": float(sig.get("recruiter_response_rate", 0) or 0),
            "interview_completion": float(sig.get("interview_completion_rate", 0) or 0),
            "offer_acceptance": offer_acc,
            "offer_acceptance_missing": 1 if offer_acc == -1 else 0,
            "github_activity": github_val,
            "github_activity_missing": 1 if github_val == -1 else 0,
            "connection_count": int(sig.get("connection_count", 0) or 0),
            "endorsements_received": int(sig.get("endorsements_received", 0) or 0),
            "views_received_30d": int(sig.get("profile_views_received_30d", 0) or 0),
            "applications_30d": int(sig.get("applications_submitted_30d", 0) or 0),
            "search_appearance_30d": int(sig.get("search_appearance_30d", 0) or 0),
            "saved_by_recruiters_30d": int(sig.get("saved_by_recruiters_30d", 0) or 0),
            "open_to_work": 1 if sig.get("open_to_work_flag") else 0,
            "verified_email": 1 if sig.get("verified_email") else 0,
            "verified_phone": 1 if sig.get("verified_phone") else 0,
            "linkedin_connected": 1 if sig.get("linkedin_connected") else 0,
            "willing_to_relocate": 1 if sig.get("willing_to_relocate") else 0,
            "advanced_skill_count": advanced_cnt,
            "expert_skill_count": expert_cnt,
            "weighted_skill_score": weighted_skill,
            "service_company_ratio": service_count / len(career) if career else 0.0,
            "skills_count": len(norm_skills),
            "career_count": len(career),
            "title_career_match_ratio": title_career_match_ratio,
        })

        text_rows.append({
            "candidate_id": cid,
            "profile_doc": profile_doc,
            "skills_doc": skills_doc_compact,
            "career_doc": career_doc_compact,
            "full_doc": full_doc,
        })

    f.close()

    np.save(outdir / "candidate_ids.npy", np.array(ids, dtype=object))
    pd.DataFrame(meta_rows).to_parquet(outdir / "candidate_metadata.parquet", index=False)
    pd.DataFrame(feat_rows).to_parquet(outdir / "candidate_preliminary_features.parquet", index=False)

    td = pd.DataFrame(text_rows)
    td["jd_text"] = jd_text
    td.to_parquet(outdir / "candidate_text.parquet", index=False)

    print(f"Preparation complete. Artifacts saved to {outdir}")

if __name__ == "__main__":
    main()
