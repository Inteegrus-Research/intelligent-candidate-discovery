#!/usr/bin/env python3
"""
Phase 3: Build 5 JD intent embeddings and save artifacts.
Outputs:
- data/artifacts/jd_intents.parquet
- data/artifacts/jd_intent_embeddings.npy
- data/artifacts/jd_intent_names.json
"""

import argparse
import json
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document

from retrieval.embedder import embed_texts
from retrieval.intent_parser import build_jd_intents, clean_text


def load_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" ".join(cells))
    return "\n".join(parts)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--jd", default="data/raw/job_description.docx")
    ap.add_argument("--outdir", default="data/artifacts")
    ap.add_argument("--model", default="all-MiniLM-L6-v2")
    args = ap.parse_args()

    jd_path = Path(args.jd)
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    jd_raw = load_docx_text(jd_path)
    jd_text = clean_text(jd_raw)

    intents = build_jd_intents(jd_text)

    intent_names = list(intents.keys())
    intent_texts = [intents[k] for k in intent_names]

    # Save text artifact
    df = pd.DataFrame({
        "intent_name": intent_names,
        "intent_text": intent_texts,
        "intent_char_len": [len(t) for t in intent_texts],
        "intent_word_len": [len(t.split()) for t in intent_texts],
    })
    df.to_parquet(outdir / "jd_intents.parquet", index=False)

    # Save embeddings
    emb = embed_texts(intent_texts, model_name=args.model, batch_size=8, normalize=True)
    np.save(outdir / "jd_intent_embeddings.npy", emb)

    # Save names for downstream code
    with open(outdir / "jd_intent_names.json", "w", encoding="utf-8") as f:
        json.dump(intent_names, f, indent=2)

    print("Phase 3 complete.")
    print(f"Saved: {outdir / 'jd_intents.parquet'}")
    print(f"Saved: {outdir / 'jd_intent_embeddings.npy'}")
    print(f"Saved: {outdir / 'jd_intent_names.json'}")
    print("Embedding shape:", emb.shape)


if __name__ == "__main__":
    main()
